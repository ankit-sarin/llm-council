"""
LLM Council - Gradio Interface

Multi-model debate system with peer review and chairman synthesis.
"""

import asyncio
import gradio as gr

from config import (
    AVAILABLE_MODELS,
    CHAIRMAN_MODEL,
    DEFAULT_ENABLED_MODELS,
    MODEL_DISPLAY_NAMES,
    MAX_CONCURRENT_VRAM_GB,
    get_display_name,
)
from council import (
    CouncilResult,
    ConsensusScore,
    ModelResponse,
    PeerReview,
    ResponseComposition,
    get_initial_responses,
    get_peer_reviews,
    get_chairman_synthesis,
    get_chairman_early_synthesis,
    stream_initial_responses,
    stream_peer_reviews,
    stream_initial_responses_live,
    stream_peer_reviews_live,
    calculate_consensus,
)

import json
import os
import re
from datetime import datetime
from pathlib import Path

# --- PII/PHI Redaction ---

# Regex patterns for common identifiers
# Order matters: more specific patterns (with context like "MRN:") should come before generic patterns
PII_PATTERNS = [
    # Email addresses
    (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'), '[EMAIL REDACTED]'),
    # Phone numbers (various formats: 123-456-7890, (123) 456-7890, 123.456.7890, +1 123 456 7890)
    (re.compile(r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'), '[PHONE REDACTED]'),
    # Medical Record Numbers - MUST come before SSN pattern (MRNs often look like SSNs but have context)
    (re.compile(r'\bMRN[:\s#]*\d{4,12}\b', re.IGNORECASE), '[MRN REDACTED]'),
    (re.compile(r'\b(?:medical\s*record|patient\s*id|chart\s*#?)[:\s#]*\d{4,12}\b', re.IGNORECASE), '[MRN REDACTED]'),
    # Dates of birth in common formats - MUST come before SSN pattern
    (re.compile(r'\b(?:DOB|date\s*of\s*birth)[:\s]*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b', re.IGNORECASE), '[DOB REDACTED]'),
    # Social Security Numbers (123-45-6789 or 123456789) - generic pattern, apply last
    (re.compile(r'\b\d{3}[-\s]?\d{2}[-\s]?\d{4}\b'), '[SSN REDACTED]'),
    # Credit card numbers (16 digits, with or without spaces/dashes)
    (re.compile(r'\b(?:\d{4}[-\s]?){3}\d{4}\b'), '[CARD REDACTED]'),
]


def redact_pii(text: str) -> str:
    """
    Redact common PII/PHI patterns from text before saving.

    Args:
        text: The text to redact

    Returns:
        Text with PII/PHI patterns replaced with redaction markers
    """
    if not text:
        return text

    redacted = text
    for pattern, replacement in PII_PATTERNS:
        redacted = pattern.sub(replacement, redacted)

    return redacted


def redact_session_data(data: dict) -> dict:
    """
    Recursively redact PII/PHI from session data before saving.

    Args:
        data: Session data dictionary

    Returns:
        Session data with PII/PHI redacted from string values
    """
    if isinstance(data, str):
        return redact_pii(data)
    elif isinstance(data, dict):
        return {k: redact_session_data(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [redact_session_data(item) for item in data]
    else:
        return data


# --- Session Storage ---

SESSIONS_DIR = Path(__file__).parent / "sessions"
SESSIONS_DIR.mkdir(exist_ok=True)


def save_session(
    question: str,
    responses: dict[str, "ModelResponse"],
    reviews: dict[str, "PeerReview"],
    synthesis: str,
    consensus: ConsensusScore,
    models: list[str],
    document_filename: str | None = None,
    document_char_count: int | None = None,
    document_preview: str | None = None,
    composition: "ResponseComposition | None" = None,
    was_stopped_early: bool = False
) -> dict:
    """
    Save a council session to JSON file.

    Args:
        question: The user's question
        responses: Dict of model responses
        reviews: Dict of peer reviews
        synthesis: Chairman's synthesis
        consensus: Consensus score object
        models: List of model IDs used
        document_filename: Optional filename if a document was uploaded
        document_char_count: Optional character count of the document
        document_preview: Optional first 1000 chars of document for reference
        composition: Optional response composition metrics
        was_stopped_early: Whether deliberation was stopped early by user
    """
    timestamp = datetime.now()
    session_id = timestamp.strftime("%Y%m%d_%H%M%S")

    session_data = {
        "id": session_id,
        "timestamp": timestamp.isoformat(),
        "question": question,
        "models": [get_display_name(m) for m in models],
        "responses": {
            get_display_name(model): {
                "content": resp.content,
                "elapsed_seconds": resp.elapsed_seconds
            }
            for model, resp in responses.items()
        },
        "reviews": {
            get_display_name(reviewer): {
                "rankings": review.rankings,
                "analysis": review.analysis
            }
            for reviewer, review in reviews.items()
        },
        "consensus": {
            "score": consensus.score,
            "level": consensus.level,
            "description": consensus.description,
            "disagreement_areas": consensus.disagreement_areas
        },
        "synthesis": synthesis,
        "was_stopped_early": was_stopped_early
    }

    # Add document info if a file was uploaded
    if document_filename:
        session_data["document"] = {
            "filename": document_filename,
            "char_count": document_char_count,
            "preview": document_preview  # First 1000 chars for reference
        }

    # Add composition info if available
    if composition:
        session_data["composition"] = {
            "council_contribution": composition.council_contribution,
            "chairman_independent": composition.chairman_independent,
            "web_search_used": composition.web_search_used,
            "web_searches_performed": composition.web_searches_performed,
            "chairman_insights": composition.chairman_insights
        }

    # Redact PII/PHI before saving
    session_data = redact_session_data(session_data)

    filepath = SESSIONS_DIR / f"session_{session_id}.json"
    with open(filepath, "w") as f:
        json.dump(session_data, f, indent=2)

    return session_data


def list_sessions() -> list[dict]:
    """List all saved sessions (newest first)."""
    sessions = []
    for filepath in sorted(SESSIONS_DIR.glob("session_*.json"), reverse=True):
        try:
            with open(filepath) as f:
                data = json.load(f)
                sessions.append({
                    "id": data["id"],
                    "timestamp": data["timestamp"],
                    "question": data["question"][:100] + ("..." if len(data["question"]) > 100 else ""),
                    "consensus_score": data["consensus"]["score"],
                    "consensus_level": data["consensus"]["level"],
                    "models": data["models"],
                    "filepath": str(filepath)
                })
        except Exception:
            continue
    return sessions


def load_session(session_id: str) -> dict | None:
    """Load a specific session by ID."""
    filepath = SESSIONS_DIR / f"session_{session_id}.json"
    if filepath.exists():
        with open(filepath) as f:
            return json.load(f)
    return None


def export_session_to_markdown(session_id: str) -> str:
    """Export a session to Markdown format."""
    session = load_session(session_id)
    if not session:
        return ""

    # Add document info if present
    document_info = ""
    if session.get("document"):
        doc = session["document"]
        document_info = f"""
**Document Analyzed:** {doc['filename']} ({doc['char_count']:,} characters)
"""

    md = f"""# Council Session: {session['id']}

**Date:** {session['timestamp']}

**Question:** {session['question']}

**Models:** {', '.join(session['models'])}{document_info}

**Consensus:** {session['consensus']['level'].upper()} ({session['consensus']['score']}/100)
{session['consensus']['description']}

---

## Individual Responses

"""
    for model, resp in session['responses'].items():
        md += f"### {model}\n*Generated in {resp['elapsed_seconds']:.1f}s*\n\n{resp['content']}\n\n---\n\n"

    md += "## Peer Reviews\n\n"
    for reviewer, review in session['reviews'].items():
        md += f"### Review by {reviewer}\n\n{review['analysis']}\n\n---\n\n"

    md += f"""## Chairman's Synthesis

{session['synthesis']}

---

*Generated by LLM Council*
"""
    return md


# --- File Upload Configuration ---
# These constants control file upload limits

MAX_FILE_SIZE_MB = 10  # Maximum file size in megabytes
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
MAX_TEXT_CHARS = 150000  # Maximum characters to extract from a file
LONG_DOC_WARNING_CHARS = 100000  # Show warning if document exceeds this
GEMMA_EXCLUSION_CHARS = 30000  # Exclude Gemma 2 (8K context) if document exceeds this


# --- File Text Extraction ---
# This section handles extracting text from uploaded files (.txt, .md, .pdf)

def extract_text_from_file(file_path: str) -> tuple[str, str | None]:
    """
    Extract text content from an uploaded file.

    Supports three file types:
    - .txt: Plain text files, read directly
    - .md: Markdown files, read directly (preserves formatting)
    - .pdf: PDF files, uses pdfplumber for better table handling

    Args:
        file_path: Path to the uploaded file

    Returns:
        Tuple of (extracted_text, error_message)
        - If successful: (text, None)
        - If failed: ("", error_message)
    """
    import os

    # Check if file exists
    if not file_path or not os.path.exists(file_path):
        return "", "File not found."

    # Get file extension (lowercase for comparison)
    file_ext = os.path.splitext(file_path)[1].lower()

    # Check if file type is supported
    supported_types = ['.txt', '.md', '.pdf', '.docx']
    if file_ext not in supported_types:
        return "", f"Please upload a .txt, .md, .pdf, or .docx file. Got: {file_ext}"

    # Check file size
    file_size = os.path.getsize(file_path)
    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        return "", f"File exceeds {MAX_FILE_SIZE_MB}MB limit. Your file: {size_mb:.1f}MB. Please use a smaller file."

    try:
        # Handle plain text files (.txt and .md)
        # These are simple - just read the contents directly
        if file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

        # Handle PDF files using pdfplumber
        # pdfplumber is better than PyMuPDF for tables
        elif file_ext == '.pdf':
            import pdfplumber

            text_parts = []  # Collect text from each page

            # Open the PDF and iterate through pages
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text from this page
                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(page_text)

                    # Also try to extract tables and format them
                    # This helps preserve table structure
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            # Convert table to text format
                            # Each row becomes a line with | separators
                            for row in table:
                                # Filter out None values and join with |
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)

            # Join all pages with double newlines
            text = "\n\n".join(text_parts)

            # Check if we got any text (might be scanned/image PDF)
            if not text.strip():
                return "", "Could not extract text from this PDF. It may be scanned or image-based."

        # Handle Word documents (.docx) using python-docx
        elif file_ext == '.docx':
            from docx import Document

            text_parts = []  # Collect text from each paragraph

            # Open the Word document
            doc = Document(file_path)

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            # Join all parts with newlines
            text = "\n\n".join(text_parts)

            if not text.strip():
                return "", "Could not extract text from this Word document."

        # Check if extraction produced any text
        if not text.strip():
            return "", "No text could be extracted from this file."

        # Truncate if text exceeds maximum
        if len(text) > MAX_TEXT_CHARS:
            text = text[:MAX_TEXT_CHARS]
            # Note: We'll show a warning in the UI, not an error

        return text, None

    except UnicodeDecodeError:
        return "", "Could not read file. It may use an unsupported character encoding."
    except Exception as e:
        return "", f"Error reading file: {str(e)}"


def get_file_info(file_path: str) -> dict:
    """
    Get information about an uploaded file for display in the UI.

    Args:
        file_path: Path to the uploaded file

    Returns:
        Dictionary with file information:
        - filename: Original filename
        - size_bytes: File size in bytes
        - size_display: Human-readable size (e.g., "1.5 MB")
        - text: Extracted text content
        - char_count: Number of characters extracted
        - preview: First 500 characters for preview
        - is_long: True if document exceeds warning threshold
        - is_truncated: True if text was truncated
        - error: Error message if extraction failed
    """
    import os

    if not file_path:
        return {"error": "No file provided"}

    # Get basic file info
    filename = os.path.basename(file_path)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

    # Format file size for display
    if file_size < 1024:
        size_display = f"{file_size} bytes"
    elif file_size < 1024 * 1024:
        size_display = f"{file_size / 1024:.1f} KB"
    else:
        size_display = f"{file_size / (1024 * 1024):.1f} MB"

    # Extract text
    text, error = extract_text_from_file(file_path)

    if error:
        return {
            "filename": filename,
            "size_bytes": file_size,
            "size_display": size_display,
            "error": error
        }

    # Calculate character count and check thresholds
    char_count = len(text)
    is_long = char_count > LONG_DOC_WARNING_CHARS
    is_truncated = char_count >= MAX_TEXT_CHARS
    exclude_gemma = char_count > GEMMA_EXCLUSION_CHARS  # Gemma 2 has 8K context limit

    # Create preview (first 500 chars)
    preview = text[:500] + ("..." if len(text) > 500 else "")

    return {
        "filename": filename,
        "size_bytes": file_size,
        "size_display": size_display,
        "text": text,
        "char_count": char_count,
        "preview": preview,
        "is_long": is_long,
        "is_truncated": is_truncated,
        "exclude_gemma": exclude_gemma,  # True if doc exceeds 30K chars
        "error": None
    }


# --- Formatting Helpers ---

def format_response_tabs(responses: dict) -> str:
    """Format individual responses as markdown with tabs simulation."""
    if not responses:
        return "*No responses yet*"

    md = ""
    for model, response in responses.items():
        display_name = get_display_name(model)
        elapsed = response.elapsed_seconds
        md += f"### {display_name}\n"
        md += f"*Generated in {elapsed:.1f}s*\n\n"
        md += f"{response.content}\n\n"
        md += "---\n\n"

    return md


def format_responses_side_by_side(partial_responses: dict, complete_responses: dict, models: list) -> str:
    """Format responses in a side-by-side grid layout with live streaming text."""
    import html as html_module

    if not partial_responses and not complete_responses:
        return "<em>Waiting for responses...</em>"

    if not models:
        return "<em>No models selected</em>"

    # Build HTML grid for side-by-side display
    num_models = len(models)
    cols = min(num_models, 3)  # Max 3 columns

    html = f'<div style="display: grid; grid-template-columns: repeat({cols}, 1fr); gap: 16px;">'

    for model in models:
        display_name = html_module.escape(get_display_name(model))
        is_complete = model in complete_responses

        # Get content
        if is_complete:
            content = complete_responses[model].content
            elapsed = complete_responses[model].elapsed_seconds
            status = f"✅ {elapsed:.1f}s"
            border_color = "#4CAF50"
        else:
            content = partial_responses.get(model, "")
            status = "⏳ generating..."
            border_color = "#FF9800"

        # Escape HTML in content - no truncation, rely on CSS scrolling
        content = html_module.escape(content) if content else ""
        display_content = content

        html += f'''
        <div style="border: 2px solid {border_color}; border-radius: 8px; padding: 12px; background: #fafafa; min-height: 200px; max-height: 600px; overflow-y: auto; overflow-x: hidden;">
            <div style="font-weight: bold; color: #333; border-bottom: 1px solid #ddd; padding-bottom: 8px; margin-bottom: 8px;">
                {display_name} <span style="font-weight: normal; color: {'#4CAF50' if is_complete else '#FF9800'};">{status}</span>
            </div>
            <div style="font-size: 0.9em; white-space: pre-wrap; word-wrap: break-word;">{display_content or "<em>Starting...</em>"}</div>
        </div>
        '''

    html += '</div>'
    return html


def format_reviews_side_by_side(partial_reviews: dict, complete_reviews: dict, models: list) -> str:
    """Format reviews in a side-by-side grid layout with live streaming text."""
    import html as html_module

    if not partial_reviews and not complete_reviews:
        return "<em>Waiting for reviews...</em>"

    if not models:
        return "<em>No reviewers</em>"

    num_models = len(models)
    cols = min(num_models, 3)

    html = f'<div style="display: grid; grid-template-columns: repeat({cols}, 1fr); gap: 16px;">'

    for model in models:
        display_name = html_module.escape(get_display_name(model))
        is_complete = model in complete_reviews

        if is_complete:
            content = complete_reviews[model].analysis
            status = "✅ complete"
            border_color = "#4CAF50"
        else:
            content = partial_reviews.get(model, "")
            status = "⏳ reviewing..."
            border_color = "#FF9800"

        # Escape HTML in content - no truncation, rely on CSS scrolling
        content = html_module.escape(content) if content else ""
        display_content = content

        html += f'''
        <div style="border: 2px solid {border_color}; border-radius: 8px; padding: 12px; background: #fff8f0; min-height: 150px; max-height: 600px; overflow-y: auto; overflow-x: hidden;">
            <div style="font-weight: bold; color: #333; border-bottom: 1px solid #ddd; padding-bottom: 8px; margin-bottom: 8px;">
                {display_name} <span style="font-weight: normal; color: {'#4CAF50' if is_complete else '#FF9800'};">{status}</span>
            </div>
            <div style="font-size: 0.9em; white-space: pre-wrap; word-wrap: break-word;">{display_content or "<em>Starting review...</em>"}</div>
        </div>
        '''

    html += '</div>'
    return html


def format_reviews(reviews: dict, responses: dict) -> str:
    """Format peer reviews as markdown."""
    if not reviews:
        return "*No reviews yet*"

    md = ""
    for reviewer, review in reviews.items():
        reviewer_name = get_display_name(reviewer)
        md += f"### Review by {reviewer_name}\n\n"

        # Show which responses they reviewed
        reviewed = [get_display_name(m) for m in responses.keys() if m != reviewer]
        md += f"*Reviewed: {', '.join(reviewed)}*\n\n"

        md += f"{review.analysis}\n\n"
        md += "---\n\n"

    return md


def format_synthesis(synthesis: str) -> str:
    """Format chairman synthesis for Tab 2 (full details)."""
    if not synthesis:
        return "*Awaiting chairman synthesis...*"

    chairman_name = get_display_name(CHAIRMAN_MODEL)
    return f"## {chairman_name}\n\n{synthesis}"


def format_consensus_meter(consensus: ConsensusScore | None) -> str:
    """Generate a visual consensus meter with score and interpretation."""
    if consensus is None:
        return "*Calculating consensus...*"

    if consensus.level == "unknown":
        return f"*{consensus.description}*"

    # Color based on level
    colors = {
        "high": ("#4CAF50", "#e8f5e9"),  # Green
        "medium": ("#FF9800", "#fff3e0"),  # Orange
        "low": ("#f44336", "#ffebee"),  # Red
    }
    fg_color, bg_color = colors.get(consensus.level, ("#9e9e9e", "#f5f5f5"))

    # Build the meter bar
    fill_pct = consensus.score
    level_emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}.get(consensus.level, "⚪")

    html = f'''
<div style="background: {bg_color}; border: 2px solid {fg_color}; border-radius: 12px; padding: 16px; margin: 8px 0;">
    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
        <span style="font-weight: bold; color: #333;">Consensus Score</span>
        <span style="font-size: 1.2em; font-weight: bold; color: {fg_color};">{level_emoji} {consensus.score:.0f}/100</span>
    </div>

    <div style="background: #e0e0e0; border-radius: 8px; height: 12px; overflow: hidden; margin-bottom: 12px;">
        <div style="background: {fg_color}; height: 100%; width: {fill_pct}%; transition: width 0.3s ease;"></div>
    </div>

    <div style="font-size: 0.95em; color: #555;">
        <strong>{consensus.level.upper()} CONSENSUS:</strong> {consensus.description}
    </div>
'''

    # Show disagreement areas if any
    if consensus.disagreement_areas:
        html += '''
    <div style="margin-top: 12px; padding-top: 12px; border-top: 1px solid #ddd;">
        <strong style="color: #333;">Key Areas of Disagreement:</strong>
        <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #666;">
'''
        for area in consensus.disagreement_areas:
            html += f'            <li>{area}</li>\n'
        html += '        </ul>\n    </div>\n'

    html += '</div>'
    return html


def calculate_agreement_summary(reviews: dict, responses: dict = None) -> str:
    """Generate agreement/disagreement summary from reviews using consensus calculation."""
    if not reviews or len(reviews) < 2:
        return "*Need at least 2 reviews to calculate consensus*"

    # Calculate actual consensus if we have responses
    if responses:
        consensus = calculate_consensus(reviews, responses)
        return format_consensus_meter(consensus)

    # Fallback if no responses provided
    num_reviewers = len(reviews)
    return f"**{num_reviewers} council members** participated in peer review."


def format_composition_meter(composition: ResponseComposition | None) -> str:
    """Generate a visual composition meter showing source breakdown."""
    if composition is None:
        return "*Calculating response composition...*"

    # Colors for each source
    council_color = "#2196F3"  # Blue for council
    chairman_color = "#9C27B0"  # Purple for chairman
    web_color = "#FF9800"  # Orange for web

    # Build the stacked bar
    council_pct = composition.council_contribution
    chairman_pct = composition.chairman_independent
    web_pct = composition.web_search_used

    html = f'''
<div style="background: #f5f5f5; border: 2px solid #9C27B0; border-radius: 12px; padding: 16px; margin: 8px 0;">
    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
        <span style="font-weight: bold; color: #333;">Response Composition</span>
    </div>

    <!-- Stacked bar -->
    <div style="background: #e0e0e0; border-radius: 8px; height: 24px; overflow: hidden; margin-bottom: 12px; display: flex;">
        <div style="background: {council_color}; height: 100%; width: {council_pct}%;" title="Council Models: {council_pct:.0f}%"></div>
        <div style="background: {chairman_color}; height: 100%; width: {chairman_pct}%;" title="Chairman Independent: {chairman_pct:.0f}%"></div>
        <div style="background: {web_color}; height: 100%; width: {web_pct}%;" title="Web Search: {web_pct:.0f}%"></div>
    </div>

    <!-- Legend -->
    <div style="display: flex; flex-wrap: wrap; gap: 16px; font-size: 0.9em;">
        <div style="display: flex; align-items: center; gap: 6px;">
            <div style="width: 16px; height: 16px; background: {council_color}; border-radius: 4px;"></div>
            <span>Council Models: <strong>{council_pct:.0f}%</strong></span>
        </div>
        <div style="display: flex; align-items: center; gap: 6px;">
            <div style="width: 16px; height: 16px; background: {chairman_color}; border-radius: 4px;"></div>
            <span>Chairman Independent: <strong>{chairman_pct:.0f}%</strong></span>
        </div>
        <div style="display: flex; align-items: center; gap: 6px;">
            <div style="width: 16px; height: 16px; background: {web_color}; border-radius: 4px;"></div>
            <span>Web Search: <strong>{web_pct:.0f}%</strong></span>
        </div>
    </div>
'''

    # Show web searches if any were performed
    if composition.web_searches_performed:
        html += '''
    <div style="margin-top: 12px; padding-top: 12px; border-top: 1px solid #ddd;">
        <strong style="color: #333;">Web Searches Performed:</strong>
        <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #666;">
'''
        for query in composition.web_searches_performed[:5]:
            import html as html_module
            html += f'            <li>{html_module.escape(query)}</li>\n'
        html += '        </ul>\n    </div>\n'

    # Show chairman insights if any
    if composition.chairman_insights:
        html += '''
    <div style="margin-top: 12px; padding-top: 12px; border-top: 1px solid #ddd;">
        <strong style="color: #333;">Chairman's Independent Insights:</strong>
        <ul style="margin: 8px 0 0 0; padding-left: 20px; color: #666;">
'''
        for insight in composition.chairman_insights[:3]:
            import html as html_module
            # Truncate long insights
            display_insight = insight[:150] + "..." if len(insight) > 150 else insight
            html += f'            <li>{html_module.escape(display_insight)}</li>\n'
        html += '        </ul>\n    </div>\n'

    html += '</div>'
    return html


# --- Main Council Runner ---

async def run_council_streaming(
    question: str,
    selected_models: list[str],
    update_queue,
    document_text: str | None = None,
    document_filename: str | None = None,
    stop_event: asyncio.Event | None = None
):
    """
    Run the council with live token streaming updates via queue.
    Shows text as models generate it in side-by-side boxes.

    Args:
        question: The user's question
        selected_models: List of model display names to use
        update_queue: Async queue for streaming UI updates
        document_text: Optional text extracted from an uploaded document
        document_filename: Optional filename of the uploaded document
        stop_event: Optional event to signal early termination
    """
    if not question.strip():
        await update_queue.put((
            "Please enter a question.",
            "",
            "No responses yet",
            "No reviews yet",
            "Awaiting synthesis...",
            "",
            "",  # Composition meter
            None  # No pending session
        ))
        return

    # Map display names back to model IDs
    name_to_id = {v: k for k, v in MODEL_DISPLAY_NAMES.items()}
    active_models = [name_to_id[name] for name in selected_models if name in name_to_id]

    if not active_models:
        await update_queue.put((
            "No council members selected! Please select at least one model.",
            "",
            "No models selected",
            "No reviews possible",
            "Cannot synthesize without responses",
            "",
            "",  # Composition meter
            None  # No pending session
        ))
        return

    model_names = [get_display_name(m) for m in active_models]
    complete_responses = {}
    complete_reviews = {}
    composition = None  # Track response composition
    was_stopped_early = False

    # Throttle updates to avoid overwhelming the UI
    import time as time_module
    last_update = [0]
    update_interval = 0.1  # Update UI every 100ms max (reduced for better proxy/tunnel compatibility)
    keepalive_interval = 5.0  # Force update every 5s to prevent proxy timeout
    last_keepalive = [0]

    # --- Stage 1: Initial Responses with live streaming ---
    await update_queue.put((
        f"**Stage 1/3:** Generating responses ({len(active_models)} models)...",
        "",
        format_responses_side_by_side({m: "" for m in active_models}, {}, active_models),
        "### ⏸️ Waiting for Stage 1...\n\n*Peer reviews will begin after all responses are complete.*",
        "### ⏸️ Waiting for responses and reviews...\n\n*Chairman will synthesize after peer reviews.*",
        "",
        "",  # Composition meter
        None  # No pending session yet
    ))

    # Token callback - update UI with partial responses
    async def on_token_stage1(partial_responses, done_responses):
        now = time_module.time()
        # Allow update if: throttle interval passed OR keepalive interval passed (prevents proxy timeout)
        needs_keepalive = now - last_keepalive[0] >= keepalive_interval
        if now - last_update[0] < update_interval and not needs_keepalive:
            return  # Throttle updates
        last_update[0] = now
        last_keepalive[0] = now

        completed = len(done_responses)
        total = len(active_models)
        status = f"**Stage 1/3:** Generating responses ({completed}/{total} complete)..."

        await update_queue.put((
            status,
            "",
            format_responses_side_by_side(partial_responses, done_responses, active_models),
            "### ⏸️ Waiting for Stage 1...\n\n*Peer reviews will begin after all responses are complete.*",
            "### ⏸️ Waiting for responses and reviews...",
            "",
            "",  # Composition meter
            None  # No pending session yet
        ))

    # Complete callback
    async def on_complete_stage1(model, response, all_done):
        complete_responses[model] = response
        completed = len(all_done)
        total = len(active_models)

        await update_queue.put((
            f"**Stage 1/3:** Generating responses ({completed}/{total} complete)...",
            "",
            format_responses_side_by_side({m: complete_responses.get(m, type('', (), {'content': ''})()).content if m in complete_responses else "" for m in active_models}, complete_responses, active_models),
            "### ⏸️ Waiting for Stage 1...",
            "### ⏸️ Waiting for responses and reviews...",
            "",
            "",  # Composition meter
            None  # No pending session yet
        ))

    # Run Stage 1 with live streaming
    # Pass document_text if a file was uploaded
    responses = await stream_initial_responses_live(
        question, active_models, on_token_stage1, on_complete_stage1,
        document_text=document_text
    )
    complete_responses = responses

    # Format timing summary
    timing_summary = "**Response Times:**\n"
    for model in active_models:
        if model in responses:
            timing_summary += f"- {get_display_name(model)}: ✅ {responses[model].elapsed_seconds:.1f}s\n"

    # Check if deliberation was stopped early
    if stop_event and stop_event.is_set():
        was_stopped_early = True
        await update_queue.put((
            "**Deliberation stopped by user.** Chairman is synthesizing with available responses...",
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            "### ⏹️ Peer reviews skipped (deliberation stopped)",
            "### 🔄 Chairman is synthesizing...",
            "",
            "",  # Composition meter - will be filled after synthesis
            None
        ))

        # Get early synthesis from chairman
        synthesis, composition = await get_chairman_early_synthesis(
            question, responses,
            document_filename=document_filename
        )

        # Prepare session data
        pending_session = {
            "question": question,
            "responses": responses,
            "reviews": {},
            "synthesis": synthesis,
            "consensus": ConsensusScore(score=0, level="unknown", description="Peer reviews not completed", disagreement_areas=[]),
            "models": active_models,
            "document_filename": document_filename,
            "document_char_count": len(document_text) if document_text else None,
            "document_preview": document_text[:1000] if document_text else None,
            "composition": composition,
            "was_stopped_early": True
        }

        # Final results with early synthesis
        final_status = f"**Council session complete (early termination)**\n\nActive members: {', '.join(model_names)}"
        composition_html = format_composition_meter(composition)

        await update_queue.put((
            final_status,
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            "### ⏹️ Peer reviews skipped (deliberation stopped)",
            format_synthesis(synthesis),
            "*Consensus not calculated - peer reviews skipped*",
            composition_html,
            pending_session
        ))
        return

    # --- Stage 2: Peer Reviews with live streaming ---
    reviewer_models = list(responses.keys())

    await update_queue.put((
        f"**Stage 2/3:** Peer review in progress ({len(responses)} reviewers)...",
        timing_summary,
        format_responses_side_by_side({}, responses, active_models),
        format_reviews_side_by_side({m: "" for m in reviewer_models}, {}, reviewer_models),
        "### ⏸️ Waiting for peer reviews...\n\n*Chairman synthesis will begin after all reviews are done.*",
        "",
        "",  # Composition meter
        None  # No pending session yet
    ))

    # Token callback for Stage 2
    async def on_token_stage2(partial_reviews, done_reviews):
        now = time_module.time()
        # Allow update if: throttle interval passed OR keepalive interval passed (prevents proxy timeout)
        needs_keepalive = now - last_keepalive[0] >= keepalive_interval
        if now - last_update[0] < update_interval and not needs_keepalive:
            return
        last_update[0] = now
        last_keepalive[0] = now

        completed = len(done_reviews)
        total = len(reviewer_models)
        consensus_html = calculate_agreement_summary(done_reviews, responses) if len(done_reviews) >= 2 else ""

        await update_queue.put((
            f"**Stage 2/3:** Peer review ({completed}/{total} complete)...",
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            format_reviews_side_by_side(partial_reviews, done_reviews, reviewer_models),
            "### ⏸️ Waiting for peer reviews...",
            consensus_html,
            "",  # Composition meter
            None  # No pending session yet
        ))

    # Complete callback for Stage 2
    async def on_complete_stage2(reviewer, review, all_done):
        complete_reviews[reviewer] = review
        completed = len(all_done)
        total = len(reviewer_models)
        consensus_html = calculate_agreement_summary(all_done, responses) if len(all_done) >= 2 else ""

        await update_queue.put((
            f"**Stage 2/3:** Peer review ({completed}/{total} complete)...",
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            format_reviews_side_by_side({m: complete_reviews.get(m, type('', (), {'analysis': ''})()).analysis if m in complete_reviews else "" for m in reviewer_models}, complete_reviews, reviewer_models),
            "### ⏸️ Waiting for peer reviews...",
            consensus_html,
            "",  # Composition meter
            None  # No pending session yet
        ))

    # Run Stage 2 with live streaming
    reviews = await stream_peer_reviews_live(question, responses, on_token_stage2, on_complete_stage2)
    complete_reviews = reviews

    # Check if deliberation was stopped during reviews
    if stop_event and stop_event.is_set():
        was_stopped_early = True
        # Use any reviews that were completed
        partial_consensus = calculate_consensus(complete_reviews, responses) if len(complete_reviews) >= 2 else None

        await update_queue.put((
            "**Deliberation stopped by user.** Chairman is synthesizing with available data...",
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            format_reviews_side_by_side({}, complete_reviews, reviewer_models),
            "### 🔄 Chairman is synthesizing...",
            format_consensus_meter(partial_consensus) if partial_consensus else "*Consensus calculation incomplete*",
            "",  # Composition meter
            None
        ))

        # Get synthesis with whatever reviews we have
        if complete_reviews:
            synthesis, composition = await get_chairman_synthesis(
                question, responses, complete_reviews, partial_consensus,
                document_filename=document_filename
            )
        else:
            synthesis, composition = await get_chairman_early_synthesis(
                question, responses,
                document_filename=document_filename
            )

        # Prepare session data
        pending_session = {
            "question": question,
            "responses": responses,
            "reviews": complete_reviews,
            "synthesis": synthesis,
            "consensus": partial_consensus or ConsensusScore(score=0, level="unknown", description="Peer reviews incomplete", disagreement_areas=[]),
            "models": active_models,
            "document_filename": document_filename,
            "document_char_count": len(document_text) if document_text else None,
            "document_preview": document_text[:1000] if document_text else None,
            "composition": composition,
            "was_stopped_early": True
        }

        final_status = f"**Council session complete (early termination)**\n\nActive members: {', '.join(model_names)}"
        composition_html = format_composition_meter(composition)

        await update_queue.put((
            final_status,
            timing_summary,
            format_responses_side_by_side({}, responses, active_models),
            format_reviews_side_by_side({}, complete_reviews, reviewer_models),
            format_synthesis(synthesis),
            format_consensus_meter(partial_consensus) if partial_consensus else "*Consensus not fully calculated*",
            composition_html,
            pending_session
        ))
        return

    # --- Stage 3: Chairman Synthesis ---
    # Calculate consensus from reviews
    consensus = calculate_consensus(reviews, responses)
    consensus_display = format_consensus_meter(consensus)

    await update_queue.put((
        "**Stage 3/3:** Chairman synthesizing final answer...\n\n*Claude Opus 4.5 is analyzing all responses and reviews...*",
        timing_summary,
        format_responses_side_by_side({}, responses, active_models),
        format_reviews_side_by_side({}, reviews, reviewer_models),
        "### 🔄 Chairman is synthesizing...\n\n*Analyzing all responses and peer reviews to create the best answer.*",
        consensus_display,
        "",  # Composition meter - will be filled after synthesis
        None  # Session data will be set after synthesis
    ))

    # Pass consensus and document info to chairman
    synthesis, composition = await get_chairman_synthesis(
        question, responses, reviews, consensus,
        document_filename=document_filename
    )

    # Prepare session data for potential saving (but don't save automatically)
    pending_session = {
        "question": question,
        "responses": responses,
        "reviews": reviews,
        "synthesis": synthesis,
        "consensus": consensus,
        "models": active_models,
        "document_filename": document_filename,
        "document_char_count": len(document_text) if document_text else None,
        "document_preview": document_text[:1000] if document_text else None,
        "composition": composition
    }

    # Final results
    final_status = f"**Council session complete!**\n\nActive members: {', '.join(model_names)}"
    composition_html = format_composition_meter(composition)

    await update_queue.put((
        final_status,
        timing_summary,
        format_responses_side_by_side({}, responses, active_models),
        format_reviews_side_by_side({}, reviews, reviewer_models),
        format_synthesis(synthesis),  # Tab 3: full synthesis
        consensus_display,  # Tab 3: consensus meter
        composition_html,  # Tab 3: composition meter
        pending_session  # Session data for manual save
    ))


def run_council_generator(
    question: str,
    selected_models: list[str],
    document_text: str | None = None,
    document_filename: str | None = None,
    stop_flag: list | None = None,
    progress=gr.Progress()
):
    """
    Generator wrapper for streaming updates to all tabs.
    Yields updates as each model completes.

    Args:
        question: The user's question
        selected_models: List of model display names to use
        document_text: Optional extracted text from uploaded file
        document_filename: Optional filename of uploaded document
        stop_flag: Optional list [bool] that can be set to True to stop deliberation
        progress: Gradio progress indicator (unused currently)
    """
    import queue
    import threading

    # Thread-safe queue for updates
    result_queue = queue.Queue()
    stop_event_holder = [None]  # To pass stop_event to async code

    def run_async():
        async def run_with_queue():
            # Create async queue
            async_queue = asyncio.Queue()

            # Create stop event
            stop_event = asyncio.Event()
            stop_event_holder[0] = stop_event

            # Task to transfer from async queue to sync queue
            async def transfer():
                while True:
                    item = await async_queue.get()
                    if item is None:
                        break
                    result_queue.put(("update", item))
                result_queue.put(("done", None))

            # Task to monitor stop_flag
            async def monitor_stop():
                while True:
                    if stop_flag and stop_flag[0]:
                        stop_event.set()
                        break
                    await asyncio.sleep(0.1)

            # Run council and transfer concurrently
            transfer_task = asyncio.create_task(transfer())
            monitor_task = asyncio.create_task(monitor_stop())

            await run_council_streaming(
                question, selected_models, async_queue,
                document_text=document_text,
                document_filename=document_filename,
                stop_event=stop_event
            )
            await async_queue.put(None)  # Signal completion
            monitor_task.cancel()
            await transfer_task

        asyncio.run(run_with_queue())

    # Start async runner in thread
    thread = threading.Thread(target=run_async)
    thread.start()

    # Yield updates as they come
    while True:
        try:
            msg_type, data = result_queue.get(timeout=600)  # 10 min timeout
            if msg_type == "done":
                break
            yield data
        except queue.Empty:
            break

    thread.join()


# --- Build the Interface ---

def create_app():
    """Create the Gradio application."""

    with gr.Blocks(title="LLM Council") as app:

        gr.Markdown(
            """
            # LLM Council
            ### Multi-model deliberation with peer review and chairman synthesis
            """,
            elem_classes="council-header"
        )

        with gr.Tabs():

            # === TAB 1: Ask the Council ===
            with gr.Tab("Ask the Council"):
                with gr.Row():
                    with gr.Column(scale=2):
                        # --- File Upload Section ---
                        gr.Markdown("#### Upload a Document (Optional)")

                        file_upload = gr.File(
                            label="Upload a file for the council to analyze",
                            file_types=[".txt", ".md", ".pdf", ".docx"],
                            file_count="single",
                        )

                        file_info_display = gr.HTML(
                            value="<em>No file uploaded. The council will answer your question directly.</em>"
                        )

                        clear_file_btn = gr.Button("🗑️ Clear File", size="sm", visible=False)

                        gr.Markdown(
                            "*Note: Graphs and images cannot be analyzed. Tables may have formatting issues in PDFs and Word docs.*",
                            elem_classes="file-note"
                        )

                        gr.Markdown("---")

                        # --- Question Input Section ---
                        question_input = gr.Textbox(
                            label="Your Question",
                            placeholder="Enter a question for the council to deliberate...",
                            lines=3,
                            elem_classes="question-input"
                        )

                        with gr.Row():
                            submit_btn = gr.Button("Submit to Council", variant="primary", size="lg")
                            stop_btn = gr.Button("⏹️ Stop Deliberation", variant="stop", size="lg", visible=False)

                        # Stop confirmation row (hidden by default)
                        with gr.Row(visible=False) as stop_confirm_row:
                            gr.Markdown("**Deliberation stopped.** What would you like to do?")
                        with gr.Row(visible=False) as stop_options_row:
                            clear_session_btn = gr.Button("🗑️ Clear Session", variant="secondary")
                            synthesize_now_btn = gr.Button("📝 Synthesize Now", variant="primary")

                    with gr.Column(scale=1):
                        # --- Model Selection ---
                        gr.Markdown("#### Select Council Members")
                        model_checkboxes = []
                        for model_id in AVAILABLE_MODELS:
                            display_name = get_display_name(model_id)
                            is_default = model_id in DEFAULT_ENABLED_MODELS
                            cb = gr.Checkbox(
                                label=display_name,
                                value=is_default,
                                elem_classes="model-checkbox"
                            )
                            model_checkboxes.append(cb)

                    with gr.Column(scale=1):
                        # --- Settings (moved from Tab 3) ---
                        gr.Markdown("#### Settings")

                        gr.Markdown(f"**Chairman:** {get_display_name(CHAIRMAN_MODEL)}")

                        gr.Markdown("---")

                        gr.Markdown("##### GPU Memory")
                        clear_gpu_btn = gr.Button("🔄 Clear GPU Memory", variant="secondary", size="sm")
                        gpu_status = gr.Markdown(value="")

                        gr.Markdown(f"*Max VRAM: {MAX_CONCURRENT_VRAM_GB}GB*")

                # Status section
                with gr.Row():
                    with gr.Column():
                        status_display = gr.Markdown(
                            value="Ready to convene the council.",
                            elem_classes="stage-indicator"
                        )
                        timing_display = gr.Markdown(value="")

            # === TAB 2: Council Deliberation ===
            with gr.Tab("Council Deliberation"):
                gr.Markdown("### Detailed view of the deliberation process")

                with gr.Accordion("Stage 1: Individual Responses", open=True, elem_classes="stage-1"):
                    responses_display = gr.HTML(value="<em>No responses yet</em>")

                with gr.Accordion("Stage 2: Peer Reviews", open=True, elem_classes="stage-2"):
                    reviews_display = gr.HTML(value="<em>No reviews yet</em>")

            # === TAB 3: Final Synthesis ===
            with gr.Tab("Final Synthesis"):
                gr.Markdown("### Chairman's Final Answer", elem_classes="final-answer-header")
                gr.Markdown(f"*Synthesized by {get_display_name(CHAIRMAN_MODEL)}*")

                synthesis_display = gr.Markdown(
                    value="*Submit a question to see the council's synthesized answer.*",
                    elem_classes="final-answer"
                )

                gr.Markdown("### Council Consensus")
                consensus_display = gr.HTML(value="<em>Submit a question to see consensus analysis</em>")

                gr.Markdown("### Response Composition")
                composition_display = gr.HTML(value="<em>Submit a question to see how the response was composed</em>")

                gr.Markdown("---")

                # Save Session button and state
                pending_session_state = gr.State(value=None)
                save_status = gr.Markdown(value="")

                save_session_btn = gr.Button(
                    "💾 Save Session to History",
                    variant="primary",
                    size="lg"
                )

                def clear_gpu_memory():
                    """Restart Ollama to clear GPU memory."""
                    import subprocess
                    try:
                        result = subprocess.run(
                            ["sudo", "systemctl", "restart", "ollama"],
                            capture_output=True,
                            text=True,
                            timeout=30
                        )
                        if result.returncode == 0:
                            return "✅ Ollama restarted successfully. GPU memory cleared."

                        subprocess.run(["pkill", "-f", "ollama"], capture_output=True, timeout=10)
                        import time
                        time.sleep(2)
                        subprocess.Popen(["ollama", "serve"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        time.sleep(3)
                        return "✅ Ollama restarted. GPU memory should be cleared."

                    except subprocess.TimeoutExpired:
                        return "⚠️ Restart timed out. Try manually: `sudo systemctl restart ollama`"
                    except Exception as e:
                        return f"❌ Failed to restart Ollama: {e}\nTry manually: `sudo systemctl restart ollama`"

                clear_gpu_btn.click(
                    fn=clear_gpu_memory,
                    outputs=[gpu_status]
                )

            # === TAB 4: History ===
            with gr.Tab("History"):
                gr.Markdown("### Past Council Sessions")
                gr.Markdown("Browse and export previous deliberations.")

                with gr.Row():
                    refresh_btn = gr.Button("🔄 Refresh", size="sm")
                    export_md_btn = gr.Button("📄 Export to Markdown", size="sm")

                session_dropdown = gr.Dropdown(
                    label="Select Session",
                    choices=[],
                    interactive=True,
                    elem_classes="session-dropdown"
                )

                with gr.Row():
                    session_info = gr.HTML(value="<em>Select a session to view details</em>")

                with gr.Accordion("Session Details", open=True):
                    session_question = gr.Markdown(value="")
                    session_consensus = gr.HTML(value="")
                    session_synthesis = gr.Markdown(value="")

                markdown_output = gr.Textbox(
                    label="Markdown Export (copy to clipboard)",
                    lines=10,
                    visible=False
                )

                # History tab callbacks
                def refresh_sessions():
                    sessions = list_sessions()
                    if not sessions:
                        return gr.update(choices=[], value=None), "<em>No sessions found. Save a session from the Final Synthesis tab.</em>"
                    choices = []
                    for s in sessions:
                        q_preview = s['question'][:50] + ("..." if len(s['question']) > 50 else "")
                        label = f"{s['timestamp'][:10]} | {s['consensus_level'].upper()} ({s['consensus_score']:.0f}) | {q_preview}"
                        choices.append((label, s['id']))
                    return gr.update(choices=choices, value=None), f"<em>Found {len(sessions)} session(s)</em>"

                def load_session_details(session_id):
                    if not session_id:
                        return "", "", ""
                    session = load_session(session_id)
                    if not session:
                        return "Session not found", "", ""

                    question_md = f"**Question:** {session['question']}\n\n**Models:** {', '.join(session['models'])}\n\n**Date:** {session['timestamp']}"

                    # Add document info if present
                    if session.get("document"):
                        doc = session["document"]
                        question_md += f"\n\n**Document:** {doc.get('filename', 'Unknown')} ({doc.get('char_count', 0):,} chars)"

                    # Add early termination note if applicable
                    if session.get("was_stopped_early"):
                        question_md += "\n\n⚠️ *Deliberation was stopped early by user*"

                    consensus = session['consensus']
                    colors = {"high": ("#4CAF50", "#e8f5e9"), "medium": ("#FF9800", "#fff3e0"), "low": ("#f44336", "#ffebee")}
                    fg, bg = colors.get(consensus['level'], ("#9e9e9e", "#f5f5f5"))
                    emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}.get(consensus['level'], "⚪")

                    consensus_html = f'''
                    <div style="background: {bg}; border: 2px solid {fg}; border-radius: 8px; padding: 12px; margin: 8px 0;">
                        <strong>{emoji} {consensus['level'].upper()} CONSENSUS</strong> ({consensus['score']}/100)<br/>
                        {consensus['description']}
                    </div>
                    '''

                    # Add composition info if available
                    if session.get("composition"):
                        comp = session["composition"]
                        council_color = "#2196F3"
                        chairman_color = "#9C27B0"
                        web_color = "#FF9800"

                        consensus_html += f'''
                        <div style="background: #f5f5f5; border: 2px solid #9C27B0; border-radius: 8px; padding: 12px; margin: 8px 0;">
                            <strong>Response Composition</strong><br/>
                            <div style="display: flex; gap: 16px; margin-top: 8px; font-size: 0.9em;">
                                <span style="color: {council_color};">● Council: {comp.get('council_contribution', 0):.0f}%</span>
                                <span style="color: {chairman_color};">● Chairman: {comp.get('chairman_independent', 0):.0f}%</span>
                                <span style="color: {web_color};">● Web: {comp.get('web_search_used', 0):.0f}%</span>
                            </div>
                        </div>
                        '''

                    synthesis_md = f"### Chairman's Synthesis\n\n{session['synthesis']}"

                    return question_md, consensus_html, synthesis_md

                def export_markdown(session_id):
                    if not session_id:
                        return gr.update(visible=False, value="")
                    md = export_session_to_markdown(session_id)
                    return gr.update(visible=True, value=md)

                refresh_btn.click(
                    fn=refresh_sessions,
                    outputs=[session_dropdown, session_info]
                )

                session_dropdown.change(
                    fn=load_session_details,
                    inputs=[session_dropdown],
                    outputs=[session_question, session_consensus, session_synthesis]
                )

                export_md_btn.click(
                    fn=export_markdown,
                    inputs=[session_dropdown],
                    outputs=[markdown_output]
                )

                # Load sessions on startup
                app.load(
                    fn=refresh_sessions,
                    outputs=[session_dropdown, session_info]
                )

        # === File Upload Handlers ===
        # These functions handle file upload, display info, and clearing

        # State to store the extracted document text
        # gr.State persists data between interactions
        document_state = gr.State(value=None)

        # State to track deliberation stop request
        # Using a list so it can be modified by reference
        stop_flag_state = gr.State(value=[False])

        # State to track if deliberation is in progress
        deliberation_in_progress = gr.State(value=False)

        def handle_file_upload(file):
            """
            Called when a user uploads a file.
            Extracts text and returns display info.

            Args:
                file: The uploaded file object from Gradio

            Returns:
                Tuple of (info_html, clear_button_visible, document_dict)
            """
            if file is None:
                # No file uploaded - reset to default state
                return (
                    "<em>No file uploaded. The council will answer your question directly.</em>",
                    gr.update(visible=False),  # Hide clear button
                    None  # Clear document state
                )

            # Get file info and extract text
            info = get_file_info(file.name)

            if info.get("error"):
                # Show error message in red
                error_html = f'''
                <div style="background: #ffebee; border: 2px solid #f44336; border-radius: 8px; padding: 12px;">
                    <strong style="color: #c62828;">❌ Error:</strong> {info["error"]}
                </div>
                '''
                return (
                    error_html,
                    gr.update(visible=True),  # Show clear button so user can try again
                    None  # No document to store
                )

            # Build success display with file info
            # Show warnings if document is long or truncated
            warnings = []
            if info.get("is_truncated"):
                warnings.append(f"⚠️ Document truncated to {MAX_TEXT_CHARS:,} characters")
            if info.get("is_long"):
                warnings.append("⚠️ Long document - some models may truncate")

            # Show Gemma exclusion note if document exceeds 30K chars
            gemma_note = ""
            if info.get("exclude_gemma"):
                gemma_note = '''
                <div style="background: #e3f2fd; border: 1px solid #2196F3; border-radius: 4px; padding: 8px; margin-top: 8px;">
                    ℹ️ <strong>Note:</strong> Gemma 2 excluded for this query due to document length (8K context limit)
                </div>
                '''

            warning_html = ""
            if warnings:
                warning_html = f'''
                <div style="background: #fff3e0; border: 1px solid #FF9800; border-radius: 4px; padding: 8px; margin-top: 8px;">
                    {" | ".join(warnings)}
                </div>
                '''

            # PHI warning banner - always shown when a document is uploaded
            phi_warning = '''
            <div style="background: #ffebee; border: 3px solid #d32f2f; border-radius: 8px; padding: 12px; margin-bottom: 12px;">
                <strong style="color: #c62828; font-size: 1.1em;">⚠️ PHI/PII WARNING</strong>
                <p style="color: #b71c1c; margin: 8px 0 0 0; font-size: 0.95em;">
                    <strong>Do not upload documents containing Protected Health Information (PHI) or sensitive personal data.</strong>
                    Document contents are sent to AI models and may be logged. If you save this session, common identifiers will be redacted but full protection cannot be guaranteed.
                </p>
            </div>
            '''

            # Format the info display
            info_html = f'''
            {phi_warning}
            <div style="background: #e8f5e9; border: 2px solid #4CAF50; border-radius: 8px; padding: 12px;">
                <div style="margin-bottom: 8px;">
                    <strong>📄 {info["filename"]}</strong> ({info["size_display"]})
                    <span style="color: #666; margin-left: 8px;">{info["char_count"]:,} characters extracted</span>
                </div>
                <div style="background: #f5f5f5; border-radius: 4px; padding: 8px; font-family: monospace; font-size: 0.85em; max-height: 150px; overflow-y: auto; white-space: pre-wrap;">
{info["preview"]}
                </div>
                {warning_html}
                {gemma_note}
            </div>
            '''

            # Store document info in state for use during council session
            document_dict = {
                "filename": info["filename"],
                "char_count": info["char_count"],
                "text": info["text"],
                "preview": info["text"][:1000],  # First 1000 chars for session saving
                "exclude_gemma": info.get("exclude_gemma", False)  # Exclude Gemma if >30K chars
            }

            return (
                info_html,
                gr.update(visible=True),  # Show clear button
                document_dict  # Store in state
            )

        def clear_file():
            """
            Called when user clicks the Clear File button.
            Resets file upload and state.
            """
            return (
                None,  # Clear file upload component
                "<em>No file uploaded. The council will answer your question directly.</em>",
                gr.update(visible=False),  # Hide clear button
                None  # Clear document state
            )

        # Connect file upload events
        file_upload.change(
            fn=handle_file_upload,
            inputs=[file_upload],
            outputs=[file_info_display, clear_file_btn, document_state]
        )

        clear_file_btn.click(
            fn=clear_file,
            outputs=[file_upload, file_info_display, clear_file_btn, document_state]
        )

        # === Main Submit Action ===
        # Wrapper to collect checkbox values, file content, and call generator
        def collect_and_run(question, document, stop_flag, *checkbox_values):
            """
            Collects all inputs and runs the council session.

            Args:
                question: The user's question
                document: Dict with document info (or None if no file)
                stop_flag: List [bool] to signal stop request
                *checkbox_values: Which models are selected
            """
            # Reset stop flag at start
            stop_flag[0] = False

            # Check if we need to exclude Gemma 2 due to document length
            exclude_gemma = document.get("exclude_gemma", False) if document else False

            # Build list of selected model display names
            selected = []
            for i, is_checked in enumerate(checkbox_values):
                if is_checked:
                    model_id = AVAILABLE_MODELS[i]
                    # Skip Gemma 2 if document exceeds 30K chars (8K context limit)
                    if exclude_gemma and "gemma2" in model_id.lower():
                        continue
                    selected.append(get_display_name(model_id))

            # Extract document text if present
            document_text = document.get("text") if document else None
            document_filename = document.get("filename") if document else None

            # Return generator results
            yield from run_council_generator(question, selected, document_text, document_filename, stop_flag)

        def on_submit_start():
            """Called when submit starts - show stop button, reset stop flag."""
            return (
                gr.update(visible=True),  # Show stop button
                gr.update(visible=False),  # Hide confirm row
                gr.update(visible=False),  # Hide options row
                [False],  # Reset stop flag
                True,  # Set deliberation in progress
            )

        def on_submit_end():
            """Called when submit ends - hide stop button."""
            return (
                gr.update(visible=False),  # Hide stop button
                False,  # Clear deliberation in progress
            )

        def on_stop_click(stop_flag):
            """Called when stop button is clicked."""
            stop_flag[0] = True
            return (
                gr.update(visible=False),  # Hide stop button
                gr.update(visible=True),  # Show confirm row
                gr.update(visible=True),  # Show options row
                stop_flag,  # Return updated flag
            )

        def on_clear_session():
            """Called when user wants to clear the session after stopping."""
            return (
                gr.update(visible=False),  # Hide confirm row
                gr.update(visible=False),  # Hide options row
                "Ready to convene the council.",  # Reset status
                "",  # Clear timing
                "<em>No responses yet</em>",  # Clear responses
                "<em>No reviews yet</em>",  # Clear reviews
                "*Submit a question to see the council's synthesized answer.*",  # Clear synthesis
                "<em>Submit a question to see consensus analysis</em>",  # Clear consensus
                "<em>Submit a question to see how the response was composed</em>",  # Clear composition
                None,  # Clear pending session
                [False],  # Reset stop flag
            )

        def on_synthesize_now():
            """Called when user wants to synthesize with current data after stopping."""
            return (
                gr.update(visible=False),  # Hide confirm row
                gr.update(visible=False),  # Hide options row
            )

        # Wire up submit button with pre/post handlers
        submit_btn.click(
            fn=on_submit_start,
            outputs=[stop_btn, stop_confirm_row, stop_options_row, stop_flag_state, deliberation_in_progress]
        ).then(
            fn=collect_and_run,
            inputs=[question_input, document_state, stop_flag_state] + model_checkboxes,
            outputs=[
                status_display,      # Tab 1: status
                timing_display,      # Tab 1: timing
                responses_display,   # Tab 2: responses
                reviews_display,     # Tab 2: reviews
                synthesis_display,   # Tab 3: chairman's synthesis
                consensus_display,   # Tab 3: consensus meter
                composition_display, # Tab 3: composition meter
                pending_session_state,  # Tab 3: session data for save
            ],
        ).then(
            fn=on_submit_end,
            outputs=[stop_btn, deliberation_in_progress]
        )

        # Wire up stop button
        stop_btn.click(
            fn=on_stop_click,
            inputs=[stop_flag_state],
            outputs=[stop_btn, stop_confirm_row, stop_options_row, stop_flag_state]
        )

        # Wire up clear session button
        clear_session_btn.click(
            fn=on_clear_session,
            outputs=[
                stop_confirm_row, stop_options_row,
                status_display, timing_display,
                responses_display, reviews_display,
                synthesis_display, consensus_display, composition_display,
                pending_session_state, stop_flag_state
            ]
        )

        # Wire up synthesize now button (just hides the options)
        synthesize_now_btn.click(
            fn=on_synthesize_now,
            outputs=[stop_confirm_row, stop_options_row]
        )

        # === Save Session Handler ===
        def handle_save_session(pending_session):
            """Save the pending session to history."""
            if pending_session is None:
                return "⚠️ No session to save. Submit a question first."

            try:
                # Extract data from pending session
                session_data = save_session(
                    question=pending_session["question"],
                    responses=pending_session["responses"],
                    reviews=pending_session["reviews"],
                    synthesis=pending_session["synthesis"],
                    consensus=pending_session["consensus"],
                    models=pending_session["models"],
                    document_filename=pending_session.get("document_filename"),
                    document_char_count=pending_session.get("document_char_count"),
                    document_preview=pending_session.get("document_preview"),
                    composition=pending_session.get("composition"),
                    was_stopped_early=pending_session.get("was_stopped_early", False)
                )
                return f"✅ Session saved! ID: {session_data['id']}"
            except Exception as e:
                return f"❌ Failed to save session: {e}"

        save_session_btn.click(
            fn=handle_save_session,
            inputs=[pending_session_state],
            outputs=[save_status]
        )

    return app


# --- Entry Point ---

if __name__ == "__main__":
    app = create_app()
    app.launch(
        server_name="0.0.0.0",
        server_port=7861,
        share=False,
        show_error=True,
        # Improve proxy/tunnel compatibility for mobile access
        # root_path handles reverse proxy path forwarding (e.g., Cloudflare Tunnel)
        root_path="",
        # Disable SSR for better streaming compatibility with reverse proxies
        ssr_mode=False,
        # Increase max_file_size for document uploads (default is often lower)
        max_file_size="15mb",
        theme=gr.themes.Soft(),
        css="""
        .council-header { text-align: center; margin-bottom: 20px; }
        .stage-indicator { font-size: 1.1em; padding: 10px; border-radius: 8px; background: #f0f0f0; }

        /* Model checkboxes styling */
        .model-checkbox {
            padding: 8px 12px;
            margin: 4px 0;
            border: 1px solid #e0e0e0;
            border-radius: 6px;
            background: #fafafa;
        }
        .model-checkbox:hover {
            background: #f0f0f0;
        }

        /* Stage 1 - Blue theme */
        .stage-1 {
            border-left: 4px solid #2196F3 !important;
            background: linear-gradient(to right, #e3f2fd, transparent) !important;
        }
        .stage-1-content {
            border-left: 2px solid #2196F3;
            padding-left: 12px;
            margin-left: 8px;
        }

        /* Stage 2 - Orange theme */
        .stage-2 {
            border-left: 4px solid #FF9800 !important;
            background: linear-gradient(to right, #fff3e0, transparent) !important;
        }
        .stage-2-content {
            border-left: 2px solid #FF9800;
            padding-left: 12px;
            margin-left: 8px;
        }

        /* Stage 3 - Green theme */
        .stage-3 {
            border-left: 4px solid #4CAF50 !important;
            background: linear-gradient(to right, #e8f5e9, transparent) !important;
        }
        .stage-3-content {
            border-left: 2px solid #4CAF50;
            padding-left: 12px;
            margin-left: 8px;
        }

        /* Chairman's Final Answer - Purple/Gold theme */
        .final-answer-header {
            color: #6A1B9A;
        }
        .final-answer {
            border: 2px solid #9C27B0;
            border-radius: 12px;
            padding: 20px;
            background: linear-gradient(135deg, #f3e5f5 0%, #fff8e1 100%);
            box-shadow: 0 4px 6px rgba(156, 39, 176, 0.1);
        }

        /* Question Input - Purple/Gold theme */
        .question-input {
            border: 2px solid #9C27B0 !important;
            border-radius: 12px !important;
            background: linear-gradient(135deg, #f3e5f5 0%, #fff8e1 100%) !important;
            box-shadow: 0 4px 6px rgba(156, 39, 176, 0.1) !important;
        }
        .question-input textarea {
            background: transparent !important;
        }
        .question-input label {
            color: #6A1B9A !important;
            font-weight: bold !important;
        }

        /* Stop button styling */
        button[variant="stop"] {
            background: #f44336 !important;
            border-color: #d32f2f !important;
            color: white !important;
        }
        button[variant="stop"]:hover {
            background: #d32f2f !important;
        }
        """,
    )
